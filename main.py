import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

HEADING_FONT_SIZE = 16
TEXT_FONT_SIZE = 10.5
TEXT_FONT_NAME = 'Consolas'
IGNORE_FILES_WITH_EXT = ['.spec.ts']  # which files to ignore
INCLUDE_FILES_WITH_EXT = ['.ts', '.tsx']  # which files to include
FOLDER_PATH = './Grocery/admin-panel'  # path to source code folder
OUTPUT = "code.docx"  # output file name


def write_code_to_doc(doc: type(Document), file_path: str):
    # formatting heading title
    title_paragraph = doc.add_paragraph(file_path.replace('\\', '/'))
    title_run = title_paragraph.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(HEADING_FONT_SIZE)

    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        code = file.read()
        if not file_path.endswith(tuple(IGNORE_FILES_WITH_EXT)):
            p = doc.add_paragraph()
            p.style.font.size = Pt(TEXT_FONT_SIZE)
            p.style.font.name = TEXT_FONT_NAME
            p.add_run(code)
    doc.add_paragraph('')
    doc.add_paragraph('')


def traverse_directory(directory, doc):
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(tuple(INCLUDE_FILES_WITH_EXT)):
                file_path = os.path.join(root, file)
                write_code_to_doc(doc, file_path)


# Main function
if __name__ == "__main__":
    target_directory = FOLDER_PATH
    output_docx_file = OUTPUT

    doc = Document()
    traverse_directory(target_directory, doc)

    doc.save(output_docx_file)
    print(f'Code documentation saved to {output_docx_file}')
